"""
Fills a stored template with field values and exports it as a .docx file.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt

from config import EXPORTS_DIR
from core.logger import get_logger

logger = get_logger(__name__)

_PLACEHOLDER_PATTERN = re.compile(r"\{([^{}]+)\}")

DEFAULT_BANGLA_FONT = "Nirmala UI"


class MissingFieldError(Exception):
    pass


@dataclass
class FilledDocument:
    title: str
    file_path: Path


def extract_placeholders(template_body: str) -> list[str]:
    seen: list[str] = []
    for match in _PLACEHOLDER_PATTERN.finditer(template_body):
        name = match.group(1).strip()
        if name not in seen:
            seen.append(name)
    return seen


def fill_template_text(template_body: str, field_values: dict[str, str]) -> str:
    required = extract_placeholders(template_body)
    missing = [name for name in required if not field_values.get(name, "").strip()]
    if missing:
        raise MissingFieldError(", ".join(missing))

    def _replace(match: re.Match[str]) -> str:
        return field_values[match.group(1).strip()]

    return _PLACEHOLDER_PATTERN.sub(_replace, template_body)


def export_to_docx(title: str, filled_text: str, font_name: str = DEFAULT_BANGLA_FONT) -> Path:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(12)

    for block in filled_text.split("\n"):
        document.add_paragraph(block)

    safe_title = re.sub(r"[\\/:*?\"<>|]", "_", title).strip() or "document"
    output_path = EXPORTS_DIR / f"{safe_title}.docx"

    counter = 1
    while output_path.exists():
        output_path = EXPORTS_DIR / f"{safe_title}_{counter}.docx"
        counter += 1

    document.save(output_path)
    logger.info("Exported document: %s", output_path)
    return output_path


def generate_document(title: str, template_body: str, field_values: dict[str, str]) -> FilledDocument:
    filled_text = fill_template_text(template_body, field_values)
    file_path = export_to_docx(title, filled_text)
    return FilledDocument(title=title, file_path=file_path)
